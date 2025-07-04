from docx import Document
from lxml import etree
import json
import json
from docx import Document
from lxml import etree

# Function to replace placeholders in the text box XML
def create_textbox_xml(data):
    return f"""
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:pict>
        <v:shape id="_x0000_s1025" type="#_x0000_t202" style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:400pt;z-index:251659264;visibility:visible;mso-wrap-style:square;mso-wrap-distance-left:9pt;mso-wrap-distance-top:9pt;mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:9pt;mso-position-horizontal:absolute;mso-position-vertical:absolute">
          <v:fill on="false"/>
          <v:stroke opacity="0"/>
          <v:textbox inset="2.88pt,2.88pt,2.88pt,2.88pt">
            <w:txbxContent>
              <w:p>
                <w:r>
                  <w:rPr>
                    <w:sz w:val="88"/> <!-- 44pt -->
                    <w:rFonts w:ascii="Jost" w:hAnsi="Jost" w:cs="Jost"/>
                  </w:rPr>
                  <w:t>{data['heading']}</w:t>
                </w:r>
              </w:p>
              <!-- Adding four empty paragraphs for spacing -->
              <w:p><w:r><w:t></w:t></w:r></w:p>
              <w:p><w:r><w:t></w:t></w:r></w:p>
              <w:p><w:r><w:t></w:t></w:r></w:p>
              <w:p><w:r><w:t></w:t></w:r></w:p>
              <w:p>
                <w:r>
                  <w:rPr>
                    <w:sz w:val="28"/> <!-- 14pt -->
                    <w:rFonts w:ascii="Jost" w:hAnsi="Jost" w:cs="Jost"/>
                  </w:rPr>
                  <w:t>Prepared By: {data['name']}</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:r>
                  <w:rPr>
                    <w:sz w:val="28"/> <!-- 14pt -->
                    <w:rFonts w:ascii="Jost" w:hAnsi="Jost" w:cs="Jost"/>
                  </w:rPr>
                  <w:t>Prepared For: {data['address']}</w:t>
                </w:r>
              </w:p>
              <w:p>
                <w:r>
                  <w:rPr>
                    <w:sz w:val="28"/> <!-- 14pt -->
                    <w:rFonts w:ascii="Jost" w:hAnsi="Jost" w:cs="Jost"/>
                  </w:rPr>
                  <w:t>Prepared On: {data['date']}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </w:r>
    """



def Intro(template_path, json_path, output_path):
  # Load the Word template and JSON file
  doc = Document(template_path)
  with open(json_path, 'r') as f:
    json_data = json.load(f)



# Create a new Document object from the existing template
  
  doc = Document(template_path)


# Generate the text box XML with replaced placeholders
  textbox_xml = create_textbox_xml(json_data)

#put the XML into json intro
  json_data['Intro'] = textbox_xml

# Save the updated JSON data back to the file (optional, if you want to keep it updated)
  with open('data.json', 'w') as file:
    json.dump(json_data, file, indent=4)

    
  txbx_content = etree.fromstring(json_data['Intro'])
  
      # Replace placeholders in the document
  for para in doc.paragraphs:
      for key, value in json_data.items():
          placeholder = '{' + key + '}'
          if placeholder in para.text:
                if key == 'Intro':
                    # Clear the placeholder text
                    para.text = ''
                    for run in para.runs:
                      run.font.name = "Jost"
                    p = para._p
                    # Append the new XML element to the paragraph
                    p.append(txbx_content)
                else:
                    para.text = para.text.replace(placeholder, str(value))



  print("Done")
# Save the updated document
  doc.save(output_path)

