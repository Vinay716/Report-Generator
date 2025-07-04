import json
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from collections import Counter
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from docx.shared import Inches, RGBColor
from docx.shared import RGBColor ,Inches ,Cm ,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import docx
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import lxml.etree  # Import lxml
import os
import tempfile
def set_cell_properties(cell, bg_color=None, font_color=None, border_color=None): # give propertyes to the tables when called
    """
    Set the background color, font color, and border color of a table cell.
    """
    tc_pr = cell._element.get_or_add_tcPr()

    # Set background color
    if bg_color:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        tc_pr.append(shd)

    # Set border color
    if border_color:
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tc_pr.append(border)

    # Set font color
    if font_color:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(font_color)
def create_custom_table(doc):
    # making of fift tabel from scrach
    # Create a table with 6 rows and 4 columns
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'  # Ensure the table has grid style

    # Set row heights for the first two rows
    for row in table.rows[:2]:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tr_height = OxmlElement('w:trHeight')
        tr_height.set(qn('w:val'), str(int(0.63 * 567)))  # Convert cm to twips
        tr_pr.append(tr_height)

    # Set column widths for the first two rows
    col_widths = [1.27, 8.43, 2.37, 2.5]
    for row in table.rows[:2]:
        for col_idx, width in enumerate(col_widths):
            if width:
                cell = row.cells[col_idx]
                cell.width = Cm(width)
                for tcW in cell._element.xpath('w:tcPr/w:tcW'):
                    tcW.set(qn('w:w'), str(int(width *532)))  # Convert cm to twips
                    tcW.set(qn('w:type'), 'dxa')

    # Merge cells in rows 3 to 6 to create one column
    for row_idx in range(2, 6):
        cell_to_merge = table.cell(row_idx, 0)
        for col_idx in range(1, 4):
            cell_to_merge.merge(table.cell(row_idx, col_idx))

    # Write headers in the first row
    headers = ["ID", "Vulnerability ", " Risk Rating", "Application/System Impacted"]
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)    # zero because it's header row
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Jost'
        # Set background color, border color, and font color of the first row
        set_cell_properties(cell, bg_color='#7030A0', font_color='FFFFFF', border_color='000000')

    # Write A, B, C, D in the 3rd to 6th rows
    entries = ["Vulnerability Description: ", "Vulnerability Evidence:  ", "Impact: ", "Recommendations: "]
    for row_idx in range(2, 6):
        cell = table.cell(row_idx, 0)
        para = cell.add_paragraph()
        run = para.add_run(entries[row_idx - 2])
        run.bold = True # Ensure the existing text is not bold (unless it was before, handle separately if needed)
        run.font.name = 'Jost'
        para.paragraph_format.left_indent = Cm(0.25)
        # Set border color of the merged cells
        set_cell_properties(cell, border_color='000000')

    # Add a page break after the table
    # doc.add_page_break()
    # doc.add_page_break()

    return table

def HTML_TO_DOC(doc,element, run=None):
    tag = "p"   
    print(len(doc.tables))
    table = doc.tables[0]
    cell = table.cell(1,0)
    para = cell.add_paragraph(style="List Bullet")
    run = para.add_run("hello wojasidkgbfajfsg")
    # for style in cell.style:
    #     print(style)
    for child in element.children:
        if isinstance(child, str) and  run != None:  # If it's a string (text)
            # print(child.parent.name)      
            # print(child.parent.name)    # add run
            if child.parent.name == "strong":
                if child.parent.parent.name == "i":
                    run.bold = True
                    run.italic = True
                else:
                    run.bold = True
                    
            elif child.parent.name == "i":
            #     print("if is working ")
                if child.parent.parent.name == "strong":
                    run.italic = True
                    run.bold = True
                else:
                    run.italic = True
        
        elif hasattr(child, 'children'):                        
            # print(child.name, "the child")  
            if child.name in ["p","li"]:
                if child.name != tag:
                    tag = child.name
                    print(child.name)
                    print(tag,"tag")
                    cell.add_paragraph()
                if child.name == "p":
                        para = cell.add_paragraph()
                elif child.name == "li":
                    if child.parent.name == "ul":
                        
                        para = cell.add_paragraph(style="List Bullet")
                        run = para.add_run("hello")
                    if child.parent.name == "ol":
                        para = cell.add_paragraph(style='List Number')
                        run = para.add_run("ok")



                    # run = para1.add_run("Division@@")    
                
                para.paragraph_format.left_indent = Cm(0.15)  # Left indentation
                para.paragraph_format.right_indent = Cm(0.15)  # Right indentation
                para.paragraph_format.space_after = Cm(0.15)  # Space after the paragraph (bottom indentation) 
                para.paragraph_format.space_before = 0
                para.paragraph_format.space_after = 0
                    
            
# 
            HTML_TO_DOC(doc,child, run)   # for terversing in the elements

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r.font.name = 'Jost'
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink
    
def hmm(soup,r, doc):
    table = doc.tables[-1]
    col = table.cell(r,0)
    for element in soup:
        if element.name == 'p':
            # Add a paragraph to the document
            para = col.add_paragraph()
            para.style = None
            for text in element.contents:
                if isinstance(text, str):
                    # Add plain text
                    run = para.add_run(text)
                    run.font.name = 'Jost'
                elif text.name == 'strong':
                    # Add bold text
                    data = text.text
                    run = para.add_run(data)
                    run.font.name = 'Jost'
                    run.font.bold = True
                elif text.name == 'i':
                    # Add italic text
                    data = text.text
                    run = para.add_run(data)
                    run.font.name = 'Jost'
                    run.font.italic = True
                elif text.name == 'a':
                    # Add link
                    href = text.get('href')
                    text = text.get_text()
                    add_hyperlink(para, text, href)
                if hasattr(text, 'contents'): 
                    for child in text.contents:
                        if isinstance(child, str):
                            continue
                        if child.name == 'strong':
                            run.font.bold = True
                            run.font.name = 'Jost'  
                        if child.name == 'i':
                            run.font.italic = True
                            run.font.name = 'Jost'
                        if child.name == 'a':
                            href = child.get('href')
                            text = child.get_text()
                            add_hyperlink(para, text, href)
        elif element.name in ['ul', 'ol']:
            # Add a list to the document
            if element.name == 'ul':
                list_style = "List Bullet"
            else:
                list_style = "List Number"
            for li in element.find_all('li'):
                para = col.add_paragraph(style=list_style)
                for text in li.contents:
                    if isinstance(text, str):
                        # Add plain text
                        run = para.add_run(text)
                        run.font.name = 'Jost'
                    elif text.name == 'strong':
                        # Add bold text
                        data = text.text
                        run = para.add_run(data)
                        run.font.bold = True
                    elif text.name == 'i':
                        # Add italic text
                        data = text.text
                        run = para.add_run(data)
                        run.font.italic = True
                    elif text.name == 'a':
                        href = text.get('href')
                        text = text.get_text()
                        add_hyperlink(para, text, href)
                    if hasattr(text, 'contents'):
                        for child in text.contents:
                            if isinstance(child, str):
                                continue
                            if child.name == 'strong':
                                run.font.bold = True
                            if child.name == 'i':
                                run.font.italic = True
                            if child.name == 'a':
                                href = child.get('href')
                                text = child.get_text()
                                add_hyperlink(para, text, href)
        else:
            para = col.add_paragraph()
            run = para.add_run(element)


# ""                
html2 = """<p>Exploiting <i>italicc<strong>BOLD what the func</strong></i><i><a href=\"https://www.google.com/\">Google.com</a>vulnerable</i><strong> <i>italic</i>XML</strong> processors by including external entities.</p>
<ul>
    <li>now <strong>BOLD <a href=\"https://www.google.com/\">Google.com</a></strong>or<a href=\"https://www.google.com/\">Google.com</a><i>italic</i></li>
    <li>or <i><strong>both</strong></i></li></ul>
<ol>
    <li><i><strong>firsst</strong></i></li>
    <li><i>secound</i></li> 
    <li>thirt</li>
</ol>
    <p>Exploiting vulnerable XML processors by including <a href=\"jesper\">jesper</a>external entities.</p><p><a href=\"https://www.google.com/\">Google.com</a></p>"""
           
html1 = """<p>Attack aimed at making a service unavailable.</p>
<ul>
    <li>hello world</li>
    <li><a href=\"www.google.com\">www.google.com</a></li>
    <li>hello<strong>bold</strong>thisso be dgoo<a href=\"google.com\"><strong>google.com</strong></a></li>
    <li><strong>bold</strong></li>
    <li><a href=\"agar.io\"><strong>agar.io</strong></a><strong>asdffd</strong></li>
    <li><strong>hello</strong><a href=\"www.google.com\"><strong>www.google.com</strong></a></li>
</ul>   
<ol>
    <li>this is number</li>
</ol>"""

html3 = "<p>helo</p> {{IMAGE_PLACEHOLDER_0_0_0_0}} <p>doejf</p>"

html4 = '''<p>Exploiting vulnerable XML processors by including external entities.</p>
<ul>
    <li>hello world</li>
    <li>bidney <strong>hotestar </strong><a href=\"https://chat.openai.com/\"><strong>https://chat.openai.com/</strong></a> this is it</li>
</ul>
<ol>
    <li>&nbsp;</li>
    <li>ok google <strong>make the font bold</strong></li>
</ol>'''

def Tables(template_path, json_path, output_path, image_paths):

    # Load the Word template and JSON file

    doc = Document(template_path)
    
    # for style in doc.styles:
    #     print(style.name)       
    
    soup = BeautifulSoup(html3, 'html.parser')
    print(soup)
    table = doc.add_table(rows=6, cols=1)
    table.style = 'Table Grid'
    soup = BeautifulSoup(html3, 'html.parser')
    create_custom_table(doc)
    hmm(soup,3, doc)
    
# Corrected loop to print the text content of each element
    for element in soup.find_all(['p', 'ul', 'ol']):
        # print(element.get_text())  # This now correctly prints the text content
        for child in element.descendants:
            pass# print(child)

    for strong_tag in soup.find_all('strong'):
        pass
        # print(strong_tag.get_text())
# Apply the hmm function to each <p>, <ul>, and <ol> element
    # for element in soup.find_all(['p', 'ul', 'ol']):
    #     print("hello")
    #     print(element.get_text())
    #     hmm(element, doc)
    # print(soup)                   #1
    # print("hello")
    # print("hello")
    # for child in soup.children:    #2
    #     print("")
    #     print(child.get_text)
    # print("hello")
    # print("hello")
    # for child in soup.children:          #3
    #     for subchild in child.children:
    #         print("")
    #         print(subchild)
    #         if subchild == "strong":
    #             print(subchild.get_text())
    # for child in soup.children:          #4
    #     for subchild in child.children:
    #         for superchil in subchild:
    #             print("")
    #             print(superchil)
    # for child in soup.children:          #5
    #     for subchild in child.children:
    #         for superchil in subchild:
    #             for omega in superchil:
    #                 print("")
    #                 print(omega)
    # table = doc.tables[0]
    # cell = table.cell(1,0)
    # cell.add_paragraph(style="List Bullet")
    # HTML_TO_DOC(doc,soup)
    # cell.add_paragraph(style="List Bullet")
    # for style in doc.styles:
    #     print(style.name)       
    with open(json_path, 'r') as f:
        data = json.load(f)
    # doc.add_paragraph(style="List Bullet")
    
                  
    # # for style in doc.styles:
    # #     print(style.name)
    # print(dir(cell.add_paragraph(style="List Paragraph")))
    print("hello")                            

    doc.save(output_path)

# def hmm(element):
    # for child in element.children:
    #     if child.name  == "p":
    #         para.add_paragraph()
    #     elif child.name  == "ul":
    #         para.add_paragraph(style = "List Bullet")
        
    #     elif child.name  == "ol":
    #         para.add_paragraph(style = "List Number")

        # No else

