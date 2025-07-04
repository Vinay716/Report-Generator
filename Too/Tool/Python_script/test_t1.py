import json
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from collections import Counter
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from PIL import Image
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
from docx.shared import Inches, RGBColor
from docx.shared import RGBColor ,Inches ,Cm ,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
import re
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import lxml.etree  # Import lxml
import os
import tempfile

def set_cell_properties(cell, bg_color=None, font_color=None, border_color=None): # give propertyes to the tables when called
    """
    Set the background color, font color, and border color of a table cell.
    """
    tc_pr = cell._element.get_or_add_tcPr()

    # Set background color
    if bg_color:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        tc_pr.append(shd)

    # Set border color
    if border_color:
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tc_pr.append(border)

    # Set font color
    if font_color:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(font_color)

def clear_table_data_but_keep_headers(table, header_rows_count=1):
    """
    Clears all content from the table except for the header row(s).
    """
    row_count = len(table.rows)
    for i in range(row_count - 1, header_rows_count - 1, -1):
        table._element.remove(table.rows[i]._element)

def List_of_in_scope_items(doc, table_data):
    # Ensure there's at least one table in the document
    if len(doc.tables) < 1:
        raise ValueError("The document does not contain any tables.")
    
    # Replace the first table
    if table_data:
        table = doc.tables[0]

        # Clear the existing table content but keep the header
        clear_table_data_but_keep_headers(table)
        
        # Add new rows for data
        for row_data in table_data:
            row = table.add_row()
            row = row.cells[0]
            row.text = row_data
            row.paragraphs[0].runs[0].font.name = 'Jost'
            set_cell_properties(row, border_color='000000')
        
        # print("New data added to the first table.")

def Testing_Accounts_Provided(doc, table_data):
    # Ensure there's at least two tables in the document
    if len(doc.tables) < 2:
        raise ValueError("The document does not contain enough tables.")
    
    # print("Number of tables in the document:", len(doc.tables))

    # selecting the secound table
    if table_data:
        table = doc.tables[1]

        # Clear the existing table content but keep headers
        clear_table_data_but_keep_headers(table)

        # Add new rows for data
        for row_data in table_data:
            row = table.add_row()
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                set_cell_properties(cell, border_color='000000')
                para = cell.paragraphs[0]
                run = para.add_run(cell_data)
                run.style.font.name = "Jost"
                # for paragraph in cell.paragraphs:   # indentation to the data in the cell
                #     print(cell)
                para.paragraph_format.left_indent = Cm(0.25)
                para.paragraph_format.right_indent = Cm(0.30)
                para.paragraph_format.space_before = Cm(0.19)
        # print("New data added to the second table.")

def Assessment_Team(doc, table_data):
    # Ensure there's at least three tables in the document
    if len(doc.tables) < 3:
        raise ValueError("The document does not contain enough tables.")
    
    # chosing the third table
    if table_data:
        table = doc.tables[2]

        # Clear the existing table content but keep headers
        clear_table_data_but_keep_headers(table, header_rows_count=1)

        print("Existing third table cleared, adding new data...")

        # Add new rows for data
        for row_data in table_data:
            row = table.add_row()
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_data
                cell.paragraphs[0].runs[0].font.name = 'Jost'
                set_cell_properties(cell, border_color='000000')
        
        print("New data added to the third table.")

def Summary_of_the_findings_tabel(doc, table_data):    # List of Vanribality table
   
    if len(doc.tables) < 4:
        raise ValueError("The document does not contain enough tables.")
    

    # Replace the fourth table data
    if table_data:
        table = doc.tables[3]

        # Clear the existing table content but keep headers
        clear_table_data_but_keep_headers(table)

        # Add new rows for data
        for row_data in table_data:
            row = table.add_row()
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_data
                cell.paragraphs[0].runs[0].font.name = 'Jost'
                set_cell_properties(cell, border_color='000000')

                # Set the left indentation of 0.19 cm
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.left_indent = Pt(0.19 * 28.3464567)  # 0.19 cm to points
        
        
        # print("New data added to the fourth table.")

def add_page_break(doc):
    # Adds a page break to the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
   

def create_custom_table(doc):
    # making of fift tabel from scrach
    # Create a table with 6 rows and 4 columns
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'  # Ensure the table has grid style

    # Set row heights for the first two rows
    for row in table.rows[:2]:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        tr_height = OxmlElement('w:trHeight')
        tr_height.set(qn('w:val'), str(int(0.63 * 567)))  # Convert cm to twips
        tr_pr.append(tr_height)

    # Set column widths for the first two rows
    col_widths = [1.27, 8.43, 2.37, 2.5]
    for row in table.rows[:2]:
        for col_idx, width in enumerate(col_widths):
            if width:
                cell = row.cells[col_idx]
                cell.width = Cm(width)
                for tcW in cell._element.xpath('w:tcPr/w:tcW'):
                    tcW.set(qn('w:w'), str(int(width *532)))  # Convert cm to twips 
                    tcW.set(qn('w:type'), 'dxa')

    # Merge cells in rows 3 to 6 to create one column
    for row_idx in range(2, 6):
        cell_to_merge = table.cell(row_idx, 0)
        for col_idx in range(1, 4):
            cell_to_merge.merge(table.cell(row_idx, col_idx))

    # Write headers in the first row
    headers = ["ID", "Vulnerability ", " Risk Rating", "Application/System Impacted"]
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)    # zero because it's header row
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Jost'
        # Set background color, border color, and font color of the first row
        set_cell_properties(cell, bg_color='#5A00A4', font_color='FFFFFF', border_color='000000')

    # Write A, B, C, D in the 3rd to 6th rows
    entries = ["Vulnerability Description: ", "Vulnerability Evidence:  ", "Impact: ", "Recommendations: "]
    for row_idx in range(2, 6):
        data = entries[row_idx - 2]
        cell = table.cell(row_idx, 0)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = 0
        run = para.add_run(data)
        run.bold = True # Ensure the existing text is not bold (unless it was before, handle separately if needed)
        run.font.name = 'Jost'
        print("ok",cell.text)
        # para.paragraph_format.left_indent = Cm(0.25)
        # Set border color of the merged cells
        set_cell_properties(cell, border_color='000000')


def Vanrability_Detailes(doc, row_2_data, rows_3_data, image_paths):
    last_index = len(row_2_data)                            # last index of the data for that no. of tables
    
    # Creation of the vanrbality detail tabel and coping it into a variable
    for i in range(0,last_index):  
        add_page_break(doc)
        create_custom_table(doc)
        tables = doc.tables
        
        # for style in doc.styles:
        #     print(style.name)
        copied_table =  tables[-1] # create a table
        # Modify the second row to change data in columns 2, 3, and 4
        if len(copied_table.rows) > 1:
            row = copied_table.rows[1]  # Takes the 2nd row to put the data in its column
            for j, new_data in zip(range(0, 4), row_2_data[i]):  # Iterate over the data of the ith element of row_2_data
                if j < len(row.cells):  # Ensure index j is within the bounds of row.cells
                    cell = row.cells[j]  # Identifying the column
                    cell.text = str(new_data)  # Ensure the data is a string
                    set_cell_properties(cell, border_color='000000')
                

        #APPENDING DATA ON ROWS
        for r,row_data in zip(range(2,6), rows_3_data[i]):   # looping in 3 to 6 rows
            
            row = copied_table.rows[r]                  # putting one row at a time 
            
            cell = row.cells[0]
            
        # Append new text with bold styling
            soup = BeautifulSoup(row_data[0], 'html.parser')
            print(soup)
            print_children(soup,r, doc)
           
            # Check if it's the row where images should be inserted
            if r == 3:  # This is the fourth row (0-based index)
                insert_images_from_paths(copied_table.rows[r], image_paths)

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r.font.name = 'Jost'
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def print_children(soup,r, doc):
    table = doc.tables[-1]
    col = table.cell(r,0)
    for element in soup:
        if element.name == 'p':
            # Add a paragraph to the document
            para = col.add_paragraph()
            para.style.font.name = "Jost"
            for text in element.contents:
                if isinstance(text, str):
                    # Add plain text
                    run = para.add_run(text)
                elif text.name == 'strong':
                    # Add bold text
                    data = text.text
                    run = para.add_run(data)
                    run.font.bold = True
                elif text.name == 'i':
                    # Add italic text
                    data = text.text
                    run = para.add_run(data)
                    run.font.italic = True
                elif text.name == 'a':
                    # Add link
                    href = text.get('href')
                    text = text.get_text()
                    add_hyperlink(para, text, href)
                if hasattr(text, 'contents'): 
                    for child in text.contents:
                        if isinstance(child, str):
                            continue
                        if child.name == 'strong':
                            run.font.bold = True
                        if child.name == 'i':
                            run.font.italic = True
                        if child.name == 'a':
                            href = child.get('href')
                            text = child.get_text()
                            add_hyperlink(para, text, href)
        elif element.name in ['ul', 'ol']:
            # Add a list to the document
            if element.name == 'ul':
                list_style = "List Bullet"
            else:
                list_style = "List Number"
            for li in element.find_all('li'):
                para = col.add_paragraph(style=list_style)
                for text in li.contents:
                    if isinstance(text, str):
                        # Add plain text
                        run = para.add_run(text)
                    elif text.name == 'strong':
                        # Add bold text
                        data = text.text
                        run = para.add_run(data)
                        run.font.bold = True
                    elif text.name == 'i':
                        # Add italic text
                        data = text.text
                        run = para.add_run(data)
                        run.font.italic = True
                    elif text.name == 'a':
                        href = text.get('href')
                        text = text.get_text()
                        add_hyperlink(para, text, href)
                    if hasattr(text, 'contents'):
                        for child in text.contents:
                            if isinstance(child, str):
                                continue
                            if child.name == 'strong':
                                run.font.bold = True
                            if child.name == 'i':
                                run.font.italic = True
                            if child.name == 'a':
                                href = child.get('href')
                                text = child.get_text()
                                add_hyperlink(para, text, href)
        else:
            para = col.add_paragraph()
            run = para.add_run(element)
                



def insert_images_from_paths(row, image_paths):  # insert image in the paragraph
    cell = row.cells[0]  # This is the specific cell
    # Finding the placeholder and replacing it which image
    for placeholder, image_path in image_paths.items():  # iterate over the placeholder one by one t
            for paragraph in cell.paragraphs:   # iterating on the paragraph one by one to see is the placeholder is in the paragraph
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, "")
                            with Image.open(image_path) as img:
                                width, height = img.size
                            
                            # Determine the image orientation (vertical or horizontal)
                            if width > height:
                                orientation = 'horizontal'
                            else:
                                orientation = 'vertical'

                            # Calculate the aspect ratio
                            aspect_ratio = width / height

                            # Set the maximum width and height for the image
                            max_width = 7  # inches
                            max_height = 6  # inches

                            # Resize the image based on its orientation and aspect ratio
                            if orientation == 'horizontal':
                                new_width = max_width
                                new_height = new_width / aspect_ratio
                            else:
                                new_height = max_height
                                new_width = new_height * aspect_ratio
                            # placeholder_paragraph = cell.add_paragraph()
                            # run = placeholder_paragraph.add_run()   
                            paragraph.add_run('\n')
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=docx.shared.Inches(new_width), height=docx.shared.Inches(new_height))
                            paragraph.add_run('\n')
                            # paragraph.add_run("")
                            paragraph.alignment = 1
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
            
def remove_paragraph(paragraph):
    """
    Remove a paragraph from the document.
    """
    # Get the parent element of the paragraph
    parent = paragraph._element.getparent()
    # Remove the paragraph element
    parent.remove(paragraph._element)
    



def create_chart_image(data):
    # Assuming data is in the format of a dictionary where keys are categories and values are counts
    categories = list(data.keys())
    counts = list(data.values())
    
    # Custom colors for each category
    colors = ['#C00000', '#FF0000', '#F79646', '#00B050', '#BDD7EE']
    
    # Create bar chart using matplotlib
    plt.figure(figsize=(8, 3))  # Adjusted figure size to match the image aspect ratio
    bar_width = 0.4  # Width of the bars
    bars = plt.bar(categories, counts, color=colors, width=bar_width)
    
    # Set the font properties for the tick labels
    plt.xticks(fontsize=12)
    plt.yticks(fontsize=12)
    
    # Set y-axis limit
    plt.ylim(0, max(counts) + 1)
    
    # Remove top and right borders and their ticks
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.tick_params(axis='both', which='both', top=False, right=False)
    
    # Remove background gridlines
    plt.gca().yaxis.grid(False)
    
    # Add data labels on top of each bar with increased fontsize (numbers in the candles)
    for bar, count in zip(bars, counts):
        yval = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2, yval + 0.5, count, ha='center', va='bottom', fontsize=12)
    
    # Save the figure to a BytesIO object
    image_stream = BytesIO()
    plt.tight_layout()
    plt.savefig(image_stream, format='png', dpi=1500)  # Increased DPI for higher quality
    image_stream.seek(0)
    plt.close()
    
    return image_stream

def remove_borders_from_cell(cell, borders):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find('.//w:tcBorders', namespaces=tc.nsmap)
    
    if tcBorders is None:
        tcBorders = lxml.etree.SubElement(tcPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')

    for border in borders:
        border_elm = tcBorders.find(f'w:{border}', tc.nsmap)
        if border_elm is None:
            border_elm = lxml.etree.SubElement(tcBorders, f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{border}')
        border_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')

        
def insert_chart_and_indicators_table_above_fourth_table(
    doc, 
    chart_width=Cm(14), 
    table_widths=[Cm(16), Cm(3)], 
    indicator_font_size=Pt(12), 
    indicator_padding_top=Pt(6), 
    indicator_padding_bottom=Pt(6), 
    indicator_padding_left=Pt(6),
    indicator_padding_right=Pt(6), 
    indicator_line_spacing=Pt(18)
):
    column_data = [row.cells[2].text for row in doc.tables[3].rows[1:]]
    # print(column_data)
    ratings = ['Critical', 'High', 'Medium', 'Low', 'Info']
    rating_counts = {rating: column_data.count(rating) for rating in ratings}
    # Create chart image in memory
    chart_image = create_chart_image(rating_counts)
    # Find the fourth table
    fourth_table = doc.tables[3]  # Fourth table is at index 3 (0-indexed)
    

    # Insert a new table with one row and two columns above the fourth table
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'

    # Set the width of the table columns
    for idx, width in enumerate(table_widths):
        tbl.cell(0, idx).width = width

    # Create the first cell (for the chart)
    cell_chart = tbl.cell(0, 0)
    paragraph_chart = cell_chart.paragraphs[0]
    run_chart = paragraph_chart.add_run()
    run_chart.add_picture(chart_image, width=chart_width)
    paragraph_chart.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create the second cell (for the indicators)
    cell_indicators = tbl.cell(0, 1)
    p_indicators = cell_indicators.paragraphs[0]
    indicators = [
        ("Critical", "#C00000"),
        ("High", "#FF0000"),
        ("Medium", "#F79646"),
        ("Low", "#00B050"),
        ("Info", "#BDD7EE")
    ]
    for label, color in indicators:
        run = p_indicators.add_run(u'\u25A0 ')  # Unicode for a solid square
        run.font.color.rgb = RGBColor.from_string(color[1:])
        run.font.size = indicator_font_size
        run.bold = True
        p_indicators.add_run(label).font.size = indicator_font_size
        p_indicators.add_run("\n")
    p_indicators.paragraph_format.space_before = indicator_padding_top
    p_indicators.paragraph_format.space_after = indicator_padding_bottom
    p_indicators.paragraph_format.left_indent = indicator_padding_left
    p_indicators.paragraph_format.right_indent = indicator_padding_right
    p_indicators.paragraph_format.line_spacing = indicator_line_spacing
    

    # Create the second row cells and remove left , right and bottem borders
    for cell in tbl.rows[1].cells:
        remove_borders_from_cell(cell, ['left', 'right', 'bottom'])


    # Move the table to just above the fourth table
    fourth_table._element.addprevious(tbl._element)

    



def Tables(template_path, json_path, output_path, image_paths):
    # Load the Word template and JSON file
    doc = Document(template_path)
    with open(json_path, 'r') as f:
        data = json.load(f)

    # Update first table
    List_of_in_scope_items(doc, data.get('first_table', {}).get('rows', []))
    
    # # Update second table
    Testing_Accounts_Provided(doc, data.get('second_table', {}).get('rows', []))
    
    # # Update third table
    Assessment_Team(doc, data.get('third_table', {}).get('rows', []))



    # # Update fourth table
    Summary_of_the_findings_tabel(doc, data.get('Vanriability_table', {}).get('fifth_table', {}).get('row_2_data', []))

    #Insertion of chart
    insert_chart_and_indicators_table_above_fourth_table(doc,  chart_width=Cm(14), table_widths=[Cm(16), Cm(3)])

    # Append data to fifth table
    Vanrability_Detailes(doc, data.get('Vanriability_table', {}).get('fifth_table', {}).get('row_2_data', []), data.get('Vanriability_table', {}).get('fifth_table', {}).get('row_3_data', []), image_paths) 
    # Vanrability_Detailes(doc, data.get('fifth_table', {}).get('row_2_data', []), data.get('fifth_table', {}).get('row_3_data', []))
    
  
    
    # Save the modified document
    doc.save(output_path)
    print("Document saved to", output_path)

# # Example usage
# # main(template_path, json_path, output_path)
